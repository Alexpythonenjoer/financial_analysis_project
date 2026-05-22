"""
Генерирует файлы различных форматов (docx, xlsx, pdf, txt, zip, rar, 7z)
с произвольным содержимым для тестирования краулера.
"""

import os
import random
import string
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF
import zipfile
import rarfile
import py7zr

# Настройки
STORAGE_DIR = './storage'
os.makedirs(STORAGE_DIR, exist_ok=True)

def random_text(length=200):
    return ''.join(random.choices(string.ascii_letters + ' ', k=length))

def generate_docx(filename):
    doc = Document()
    doc.add_heading('Тестовый документ', level=1)
    doc.add_paragraph(random_text(500))
    doc.save(filename)

def generate_xlsx(filename):
    wb = Workbook()
    ws = wb.active
    ws.title = "Данные"
    ws.append(["Имя", "Возраст", "Город"])
    for _ in range(10):
        ws.append([random_text(5), random.randint(18, 60), random_text(8)])
    wb.save(filename)

def generate_pdf(filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Тестовый PDF документ", ln=True, align='C')
    pdf.multi_cell(0, 10, txt=random_text(1000))
    pdf.output(filename)

def generate_txt(filename):
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(random_text(1000))

def create_archive(archive_path, files_to_archive):
    if archive_path.endswith('.zip'):
        with zipfile.ZipFile(archive_path, 'w') as z:
            for f in files_to_archive:
                z.write(f, arcname=os.path.basename(f))
    elif archive_path.endswith('.rar'):
        # Для RAR нужен установленный rarfile и утилита unrar
        # Создадим через rarfile (требует установки rar)
        # Если нет возможности, можно пропустить или использовать zip
        print("RAR создание пропущено (требуется утилита). Создан ZIP вместо него.")
        create_archive(archive_path.replace('.rar', '.zip'), files_to_archive)
    elif archive_path.endswith('.7z'):
        with py7zr.SevenZipFile(archive_path, 'w') as sz:
            for f in files_to_archive:
                sz.write(f, arcname=os.path.basename(f))

# Генерация одиночных файлов
print("Генерация файлов...")
generate_docx(os.path.join(STORAGE_DIR, 'sample1.docx'))
generate_xlsx(os.path.join(STORAGE_DIR, 'data.xlsx'))
generate_pdf(os.path.join(STORAGE_DIR, 'report.pdf'))
generate_txt(os.path.join(STORAGE_DIR, 'note.txt'))

# Создадим несколько дополнительных файлов для вложенности
subdir = os.path.join(STORAGE_DIR, 'subfolder')
os.makedirs(subdir, exist_ok=True)
generate_docx(os.path.join(subdir, 'sub_doc.docx'))
generate_txt(os.path.join(subdir, 'sub_note.txt'))

# Создание архивов
files_to_zip = [
    os.path.join(STORAGE_DIR, 'sample1.docx'),
    os.path.join(STORAGE_DIR, 'note.txt')
]
create_archive(os.path.join(STORAGE_DIR, 'archive.zip'), files_to_zip)

# Архив с вложенными файлами внутри
inner_files = [
    os.path.join(subdir, 'sub_doc.docx'),
    os.path.join(subdir, 'sub_note.txt')
]
create_archive(os.path.join(STORAGE_DIR, 'nested.zip'), inner_files)

print(f"Тестовые файлы созданы в папке {STORAGE_DIR}")